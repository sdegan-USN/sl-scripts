import pandas as pd
import docx
import pdb
import csv
from docx.oxml import OxmlElement


# open the Word document
doc = docx.Document('output.docx')

# create an empty list to store the tables
tables = []

# loop through each table in the document
for i, table in enumerate(doc.tables):

    # check if at least one cell in the header row has white text
    has_white_text = any(cell.text.strip() == '' or cell.paragraphs[0].runs[0].font.color.rgb == docx.shared.RGBColor(255, 255, 255) for row in table.rows for cell in row.cells)

    # if the header row has at least one cell with white text, add the table to the list
    if has_white_text:
        print("Table index "+str(i)+" has been selected.")
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace('\n', '').replace('\r', '')
                if len(row_data)==0:
                    row_data.append(str(cell_text))
                if len(row_data)>0 and str(cell_text) != row_data[-1]:
                    row_data.append(str(cell_text))
                #if i==14:
                #    print("Added row")
                #    pdb.set_trace()
            table_data.append(row_data)
            #tables.append(pd.DataFrame(row_data))
        #tables.append(pd.DataFrame(table_data[1:], columns=table_data[0]))
        #tables.append(table_data)
        tables.append(pd.DataFrame(table_data, columns=None))
        print("Index of table is "+str(i))
        print(" ")
        print(tables)

#with open('temp.csv', 'w', newline='') as csvfile:
#    csvwriter = csv.writer(csvfile)
#
#    for sublist in tables:
#        for row in sublist:
#            csvwriter.writerow(row)

# create a new Excel file and write the tables to separate sheets
row_num = 0
with pd.ExcelWriter('output.xlsx') as writer:
    workbook = writer.book  # Get a handle to the workbook
    worksheet = workbook.add_worksheet('Sheet1')  # Add a new worksheet
    bold_format = workbook.add_format({'bold': True})  # Define a bold format for the titles

    for i, table in enumerate(tables):
        title = f'Title {i+1}.'
        worksheet.write(row_num, 0, title, bold_format)  # Write the title
        #table.to_excel(writer, sheet_name=f'Table {i+1}', index=False, header=False)
        table.to_excel(writer, sheet_name=f'Sheet1', startrow=row_num+1, index=False, header=False)
        row_num += len(table.index) +2

