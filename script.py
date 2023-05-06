from docx import Document

# Load the Word document
document = Document('test.docx')

# Ask the user for the number of tables in the document
num_tables = int(input(f"Found {len(document.tables)} tables in the document. How many tables do you want to export? "))

# Generate the new Python script
new_script = f"""from docx import Document
from openpyxl import Workbook

# Load the Word document
document = Document('test.docx')

# Create a new Excel workbook
workbook = Workbook()
worksheet = workbook.active

"""

for i, table in enumerate(document.tables):
    if i >= num_tables:
        break
    title = f"Table {i + 1}"
    new_script += f"""
# {title}
worksheet.append(['{title}'])
"""
    for row in table.rows:
        new_row = []
        for cell in row.cells:
            text = ''
            for paragraph in cell._element.xpath('.//w:p'):
                for run in paragraph.xpath('.//w:r'):
                    text += run.text
            new_row.append(text)
        new_script += f"worksheet.append({new_row})\n"
    new_script += "worksheet.append([])\n"

new_script += "# Save the Excel workbook\nworkbook.save('tables.xlsx')"
    
# Save the new Python script to a file
with open('new_script.py', 'w') as f:
    f.write(new_script)

